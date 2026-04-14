from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches

def add_native_toc(paragraph, instruction):
    """
    Inserts a Word field (like a TOC) into the given paragraph.
    `instruction` examples:
    - Table of Contents: 'TOC \\o "1-3" \\h \\z \\u'
    - List of Figures: 'TOC \\h \\z \\c "Figure"'
    - List of Tables: 'TOC \\h \\z \\c "Table"'
    """
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar)

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = instruction
    run._r.append(instrText)

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    run._r.append(fldChar2)

    # We generally need some placeholder text between separate and end for Word to update it inside the UI
    # but an empty updateable field is fine. Word will prompt to update it.
    
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar3)

def re_insert_native_frontmatter():
    doc_path = r'c:\Users\Harsha\Documents\muse-ai\muse-ai\Muse_AI_College_Project_Report_Final_Structured.docx'
    doc = Document(doc_path)
    
    # Let's remove the raw text frontmatter we just inserted.
    # We will look for "TABLE OF CONTENTS" up to "1. ABSTRACT" and clear it out.
    toc_start_idx = -1
    abstract_idx = -1
    
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip().upper()
        if text == 'TABLE OF CONTENTS':
            toc_start_idx = i
        if text == '1. ABSTRACT':
            abstract_idx = i
            break
            
    if toc_start_idx != -1 and abstract_idx != -1 and abstract_idx > toc_start_idx:
        # Delete old frontmatter
        for i in range(abstract_idx - 1, toc_start_idx - 1, -1):
            p = doc.paragraphs[i]
            p._element.getparent().remove(p._element)
            
    # Locate 1. ABSTRACT again
    abstract_idx = -1
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == '1. ABSTRACT':
            abstract_idx = i
            break
            
    if abstract_idx == -1:
        print("Error: Could not find '1. ABSTRACT'.")
        # Just grab something reasonable
        abstract_idx = len(doc.paragraphs) // 2
        
    abstract_para = doc.paragraphs[abstract_idx]
    
    def insert_heading(text):
        p = abstract_para.insert_paragraph_before(text)
        p.style = 'Heading 1'
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return p

    def insert_pb():
        abstract_para.insert_paragraph_before().add_run().add_break()
        
    # --- NATIVE TABLE OF CONTENTS ---
    insert_heading("TABLE OF CONTENTS")
    p_toc = abstract_para.insert_paragraph_before()
    add_native_toc(p_toc, 'TOC \\o "1-3" \\h \\z \\u')
    # Note for user
    desc = abstract_para.insert_paragraph_before("(Right-click the area above and select 'Update Field' in MS Word to render the interactive Table of Contents)")
    desc.style.font.italic = True
    insert_pb()
    
    # --- NATIVE LIST OF FIGURES ---
    insert_heading("LIST OF FIGURES")
    p_lof = abstract_para.insert_paragraph_before()
    add_native_toc(p_lof, 'TOC \\h \\z \\c "Figure"')
    desc2 = abstract_para.insert_paragraph_before("(Right-click the area above and select 'Update Field' in MS Word to render the interactive List of Figures)")
    desc2.style.font.italic = True
    insert_pb()
    
    # --- NATIVE LIST OF TABLES ---
    insert_heading("LIST OF TABLES")
    p_lot = abstract_para.insert_paragraph_before()
    add_native_toc(p_lot, 'TOC \\h \\z \\c "Table"')
    desc3 = abstract_para.insert_paragraph_before("(Right-click the area above and select 'Update Field' in MS Word to render the interactive List of Tables)")
    desc3.style.font.italic = True
    insert_pb()
    
    # --- ACRONYMS AND ABBREVIATIONS ---
    insert_heading("ACRONYMS AND ABBREVIATIONS")
    acronyms = {
        "API": "Application Programming Interface",
        "AI": "Artificial Intelligence",
        "UI": "User Interface",
        "UX": "User Experience",
        "CTA": "Call to Action",
        "JSON": "JavaScript Object Notation",
        "HTML": "HyperText Markup Language",
        "CSS": "Cascading Style Sheets",
        "REST": "Representational State Transfer",
        "HTTPS": "Hypertext Transfer Protocol Secure",
        "DOM": "Document Object Model",
        "QA": "Quality Assurance",
        "SPA": "Single Page Application",
        "LLM": "Large Language Model"
    }
    
    for abbr, full_form in acronyms.items():
        p = abstract_para.insert_paragraph_before()
        p.add_run(f"{abbr}").bold = True
        p.add_run(f"\t- {full_form}")
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls
        tab_stops = parse_xml(r'<w:tabs %s><w:tab w:val="left" w:pos="1440"/></w:tabs>' % nsdecls('w'))
        p.paragraph_format._element.get_or_add_pPr().append(tab_stops)
    insert_pb()

    doc.save(doc_path)
    print("Successfully built Native Word Fields for hyperlinked TOC, TOF, and TOT!")

if __name__ == '__main__':
    re_insert_native_frontmatter()