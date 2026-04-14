from docx import Document
from docx.shared import Inches
from PIL import Image, ImageDraw
import os

def create_api_diagram(filename):
    img = Image.new('RGB', (800, 400), color=(255, 255, 255))
    d = ImageDraw.Draw(img)
    
    # Title
    d.text((250, 20), "Custom Hero & Body Generation API Architecture", fill=(0, 0, 0))
    
    # Client App
    d.rectangle([50, 150, 200, 250], outline=(0, 0, 0), width=2, fill=(230, 240, 255))
    d.text((90, 190), "React Client\n(User Inputs)", fill=(0, 0, 0))
    
    # Arrow to API
    d.line([200, 200, 300, 200], fill=(0,0,0), width=2)
    d.polygon([(290, 195), (300, 200), (290, 205)], fill=(0,0,0))
    d.text((210, 180), "POST JSON", fill=(50,50,50))
    
    # API Node
    d.rectangle([300, 100, 450, 300], outline=(0,0,0), width=2, fill=(255, 230, 230))
    d.text((320, 180), "Custom Content\nGenerator API\n(Hero & Body)", fill=(0,0,0))
    
    # Output Arrows
    d.line([450, 150, 550, 150], fill=(0,0,0), width=2)
    d.polygon([(540, 145), (550, 150), (540, 155)], fill=(0,0,0))
    
    d.line([450, 250, 550, 250], fill=(0,0,0), width=2)
    d.polygon([(540, 245), (550, 250), (540, 255)], fill=(0,0,0))
    
    # Results
    d.rectangle([550, 80, 750, 180], outline=(0,0,0), width=2, fill=(230, 255, 230))
    d.text((580, 120), "Dynamic Hero Banner\n(Title, Subtitle, CTAs)", fill=(0,0,0))
    
    d.rectangle([550, 220, 750, 320], outline=(0,0,0), width=2, fill=(230, 255, 230))
    d.text((580, 260), "Structured Body\n(Sections, Layout)", fill=(0,0,0))
    
    img.save(filename)

def append_diagram(doc_path):
    try:
        os.makedirs(r'c:\Users\Harsha\Documents\muse-ai\muse-ai\backend\temp_diagrams', exist_ok=True)
        img_path = r'c:\Users\Harsha\Documents\muse-ai\muse-ai\backend\temp_diagrams\api_architecture.png'
        
        create_api_diagram(img_path)
        
        doc = Document(doc_path)
        
        # Add the picture right after our recent text
        doc.add_paragraph('Figure 6.1: Custom Hero and Body Generation API Architecture Pipeline', style='Caption')
        try:
            doc.add_picture(img_path, width=Inches(6.0))
            print("Successfully updated DOCX with diagram!")
        except Exception as pic_err:
            print("Failed to add picture:", pic_err)
            
        doc.save(doc_path)
    except Exception as e:
        print(f"Error appending diagram: {e}")

if __name__ == '__main__':
    append_diagram(r'c:\Users\Harsha\Documents\muse-ai\muse-ai\Muse_AI_College_Project_Report_Final_Structured.docx')