from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_native_toc(paragraph, instruction):
    """
    Inserts a Word field (like a TOC) into the given paragraph.
    `instruction` examples:
    - List of Figures: 'TOC \\h \\z \\c "Figure"'
    - List of Tables: 'TOC \\h \\z \\c "Table"'
    """
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar)

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = instruction
    run._r.append(instrText)

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    run._r.append(fldChar2)

    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar3)

def re_insert_tof_tot():
    doc_path = r'c:\Users\Harsha\Documents\muse-ai\muse-ai\Muse_AI_College_Project_Report_Final_Structured.docx'
    doc = Document(doc_path)
    
    # Locate 1. ABSTRACT so we insert everything right before it.
    abstract_idx = -1
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == '1. ABSTRACT':
            abstract_idx = i
            break
            
    if abstract_idx == -1:
        print("Error: Could not find '1. ABSTRACT'.")
        # Just grab something reasonable
        abstract_idx = len(doc.paragraphs) // 2
        
    abstract_para = doc.paragraphs[abstract_idx]
    
    def insert_heading(text):
        p = abstract_para.insert_paragraph_before(text)
        p.style = 'Heading 1'
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return p

    def insert_pb():
        abstract_para.insert_paragraph_before().add_run().add_break()
    
    # Check if 'LIST OF FIGURES' already exists and skip if so.
    found_tof = False
    for p in doc.paragraphs[:abstract_idx]:
        if 'LIST OF FIGURES' in p.text.upper() or 'TABLE OF FIGURES' in p.text.upper():
            found_tof = True
            break
            
    if not found_tof:
        # --- NATIVE LIST OF FIGURES ---
        insert_heading("LIST OF FIGURES")
        p_lof = abstract_para.insert_paragraph_before()
        add_native_toc(p_lof, 'TOC \\h \\z \\c "Figure"')
        desc2 = abstract_para.insert_paragraph_before("(Right-click the area above and select 'Update Field' in MS Word to render the interactive List of Figures)")
        desc2.style.font.italic = True
        insert_pb()
        
        # --- NATIVE LIST OF TABLES ---
        insert_heading("LIST OF TABLES")
        p_lot = abstract_para.insert_paragraph_before()
        add_native_toc(p_lot, 'TOC \\h \\z \\c "Table"')
        desc3 = abstract_para.insert_paragraph_before("(Right-click the area above and select 'Update Field' in MS Word to render the interactive List of Tables)")
        desc3.style.font.italic = True
        insert_pb()

    doc.save(doc_path)
    
    if found_tof:
        print("The Table of Figures and Tables were already generated natively!")
    else:
        print("Successfully generated Native Word Fields for hyperlinked TOF and TOT!")

if __name__ == '__main__':
    re_insert_tof_tot()