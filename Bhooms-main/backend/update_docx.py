from docx import Document
import sys

def update_doc(file_path):
    try:
        doc = Document(file_path)
        doc.add_heading('6. Recent Updates', level=1)
        doc.add_heading('6.1 Custom Hero and Body Generation API', level=2)
        doc.add_paragraph(
            "As part of the continuous evolution of Muse-AI, a new dedicated API has been integrated to specifically handle the dynamic generation of 'custom hero' and 'body' sections for websites. "
            "This API provides an enhanced level of customization, allowing users to inject specific contextual information, preferences, and dynamic assets directly into the hero banner (the primary focal point of the website) as well as the main body content."
        )
        p = doc.add_paragraph()
        p.add_run("Key Features of the New API:\n").bold = True
        p.add_run("• Dynamic Hero Section: The API interprets user inputs to structure the hero title, subtitle, and primary calls-to-action (CTAs) with highly engaging, context-aware content.\n")
        p.add_run("• Body Content Structuring: It breaks down the main body into logically flowing sections (about, services, features) dynamically tailored to the specific business niche.\n")
        p.add_run("• Enhanced Customization: Users are no longer restricted to static templates; the API algorithmically adjusts typography, spacing, and asset positioning to accommodate real-time customized hero and body layouts.")
        
        doc.save(file_path)
        print("Successfully updated the DOCX file!")
    except Exception as e:
        print(f"Error updating doc: {e}")

if __name__ == '__main__':
    update_doc(r'c:\Users\Harsha\Documents\muse-ai\muse-ai\Muse_AI_College_Project_Report_Final_Structured.docx')