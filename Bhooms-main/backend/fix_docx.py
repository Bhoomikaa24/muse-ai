from docx import Document
from docx.shared import Inches
from PIL import Image, ImageDraw
import os

def create_flowchart(filename):
    img = Image.new('RGB', (800, 600), color=(255, 255, 255))
    d = ImageDraw.Draw(img)
    # Background and Title
    d.text((300, 20), "Muse AI Application Flowchart", fill=(0,0,0))
    
    # Nodes
    d.rectangle([300, 60, 500, 120], outline=(0,0,0), width=2, fill=(240, 240, 240))
    d.text((340, 80), "User Opens App", fill=(0,0,0))
    
    d.line([400, 120, 400, 160], fill=(0,0,0), width=2)
    
    d.rectangle([300, 160, 500, 220], outline=(0,0,0), width=2, fill=(230, 240, 255))
    d.text((310, 180), "Enter Business Info\n(Name, Contact, Services)", fill=(0,0,0))
    
    d.line([400, 220, 400, 260], fill=(0,0,0), width=2)
    
    d.rectangle([300, 260, 500, 320], outline=(0,0,0), width=2, fill=(255, 230, 230))
    d.text((330, 275), "Custom API Generation\n(Hero & Body Sections)", fill=(0,0,0))
    
    d.line([400, 320, 400, 360], fill=(0,0,0), width=2)
    
    d.rectangle([300, 360, 500, 420], outline=(0,0,0), width=2, fill=(230, 255, 230))
    d.text((320, 380), "Render Generated UI\n(Live Preview)", fill=(0,0,0))
    
    d.line([400, 420, 400, 460], fill=(0,0,0), width=2)
    
    d.rectangle([300, 460, 500, 520], outline=(0,0,0), width=2, fill=(240, 240, 255))
    d.text((350, 480), "Export/Publish", fill=(0,0,0))
    
    img.save(filename)

def modify_doc():
    doc_path = r'c:\Users\Harsha\Documents\muse-ai\muse-ai\Muse_AI_College_Project_Report_Final_Structured.docx'
    doc = Document(doc_path)
    
    # 1. Remove the "6. Recent Updates" appended recently
    delete_from_index = -1
    for i, p in enumerate(doc.paragraphs):
        if p.text == '6. Recent Updates':
            delete_from_index = i
            break
            
    if delete_from_index != -1:
        # Delete all paragraphs from this index to the end
        for i in range(len(doc.paragraphs) - 1, delete_from_index - 1, -1):
            p = doc.paragraphs[i]
            p._element.getparent().remove(p._element)
    
    # 2. Add API details and flowchart in System Architecture/System Flow chapter
    # Seek paragraphs where 'System Architecture' or 'System Flow'
    # Actually, let's insert it inside '7. System Architecture Overview' or similar.
    # As fallback we just append at the end of the 9th chapter or equivalent text
    arch_index = -1
    for i, p in enumerate(doc.paragraphs):
        text_lower = p.text.lower().strip()
        # Find chapter starting like 'system design' or 'system architecture'
        if 'system architecture overview' in text_lower or 'system design &' in text_lower:
            arch_index = i + 1
            break
            
    if arch_index == -1:
        print("Could not find System Architecture chapter precisely. Instead, inserting before the first 'Conclusion' or similar...")
        for i, p in enumerate(doc.paragraphs):
            if 'conclusion' in p.text.lower().strip() and len(p.text) < 25:
                arch_index = i
                break
    if arch_index == -1:
        arch_index = len(doc.paragraphs) # Append at the end if not found
    
    os.makedirs(r'c:\Users\Harsha\Documents\muse-ai\muse-ai\backend\temp_diagrams', exist_ok=True)
    fc_path = r'c:\Users\Harsha\Documents\muse-ai\muse-ai\backend\temp_diagrams\flowchart.png'
    create_flowchart(fc_path)
        
    doc.paragraphs[arch_index].insert_paragraph_before('Application Flow and Custom Hero & Body API Generation').style = 'Heading 3'
    
    doc.paragraphs[arch_index].insert_paragraph_before(
        "The application integrates a dedicated Custom Hero & Body Generation API which operates in real-time as users progress through the form. "
        "Upon capturing business information, instead of limiting the output to pre-defined static templates, the API generates a customized Hero Section (comprising uniquely tailored titles, subtitles, and calls-to-action) and structures main Body Sections tailored distinctively for the business niche. "
        "This flow creates logically spaced elements, fluid sections, and responsive UI components adapted for web rendering."
    )
    
    doc.paragraphs[arch_index].insert_paragraph_before('Figure: Muse AI Core Application and Custom Layout Pipeline', style='Caption')
    
    p_img = doc.paragraphs[arch_index].insert_paragraph_before()
    run = p_img.add_run()
    run.add_picture(fc_path, width=Inches(5.0))
    p_img.alignment = 1 # Center align

    doc.save(doc_path)
    print("Document successfully updated and flowchart inserted appropriately!")

if __name__ == '__main__':
    modify_doc()