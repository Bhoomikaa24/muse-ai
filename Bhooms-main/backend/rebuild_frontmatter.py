from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches

def extract_toc(doc):
    toc = []
    figures = []
    tables = []
    
    # 1. Gather all headings, figures, tables
    start_collecting = False
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
            
        if text.startswith('1. ABSTRACT'):
            start_collecting = True
            
        if start_collecting:
            style_name = p.style.name
            if style_name == 'Heading 1':
                toc.append(('h1', text))
            elif style_name == 'Heading 2':
                toc.append(('h2', text))
            elif style_name == 'Heading 3':
                toc.append(('h3', text))
                
            if text.lower().startswith('figure') or text.startswith('• Figure:'):
                figures.append(text.replace('• ', ''))
            elif text.lower().startswith('table '):
                tables.append(text)
                
    acronyms = {
        "API": "Application Programming Interface",
        "AI": "Artificial Intelligence",
        "UI": "User Interface",
        "UX": "User Experience",
        "CTA": "Call to Action",
        "JSON": "JavaScript Object Notation",
        "HTML": "HyperText Markup Language",
        "CSS": "Cascading Style Sheets",
        "REST": "Representational State Transfer",
        "HTTPS": "Hypertext Transfer Protocol Secure",
        "DOM": "Document Object Model",
        "QA": "Quality Assurance",
        "DB": "Database",
        "SPA": "Single Page Application",
        "LLM": "Large Language Model"
    }
    
    return toc, figures, tables, acronyms

def rebuild_frontmatter():
    doc_path = r'c:\Users\Harsha\Documents\muse-ai\muse-ai\Muse_AI_College_Project_Report_Final_Structured.docx'
    doc = Document(doc_path)
    
    toc_data, figures_data, tables_data, acronyms_data = extract_toc(doc)
    
    # Tables are often not captioned properly, if empty let's find tables in Doc? 
    # Let's add some default if none are found but tables actually exist mechanically
    if not tables_data and len(doc.tables) > 0:
        for i in range(len(doc.tables)):
            tables_data.append(f"Table {i+1}: Data Structure Reference {i+1}")

    # 2. Find deletion bounds
    toc_start_idx = -1
    abstract_idx = -1
    
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip().upper()
        if text == 'TABLE OF CONTENTS':
            toc_start_idx = i
        if text == '1. ABSTRACT':
            abstract_idx = i
            break
            
    if toc_start_idx != -1 and abstract_idx != -1 and abstract_idx > toc_start_idx:
        # Delete old frontmatter
        for i in range(abstract_idx - 1, toc_start_idx - 1, -1):
            p = doc.paragraphs[i]
            p._element.getparent().remove(p._element)

    # 3. Create the new front matter *before* 1. ABSTRACT
    # Since we deleted things, abstract_idx might have shifted. 
    # Let's find 1. ABSTRACT again.
    abstract_idx = -1
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == '1. ABSTRACT':
            abstract_idx = i
            break

    if abstract_idx == -1:
        print("Error: Could not find '1. ABSTRACT' to insert before.")
        return

    abstract_para = doc.paragraphs[abstract_idx]
    
    def insert_heading(text):
        p = abstract_para.insert_paragraph_before(text)
        p.style = 'Heading 1'
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return p

    def insert_item(text, indent=0):
        p = abstract_para.insert_paragraph_before(text)
        p.paragraph_format.left_indent = Inches(indent)
    
    def insert_pb():
        abstract_para.insert_paragraph_before().add_run().add_break()
    
    # --- TABLE OF CONTENTS ---
    insert_heading("TABLE OF CONTENTS")
    for level, text in toc_data:
        if level == 'h1':
            insert_item(text, indent=0)
        elif level == 'h2':
            insert_item(text, indent=0.3)
        elif level == 'h3':
            insert_item(text, indent=0.6)
    insert_pb()
    
    # --- LIST OF FIGURES ---
    insert_heading("LIST OF FIGURES")
    if figures_data:
        for fig in figures_data:
            insert_item(fig)
    else:
        insert_item("No figures present.")
    insert_pb()
        
    # --- LIST OF TABLES ---
    insert_heading("LIST OF TABLES")
    if tables_data:
        for tbl in tables_data:
            insert_item(tbl)
    else:
        insert_item("No tables present.")
    insert_pb()
        
    # --- ACRONYMS AND ABBREVIATIONS ---
    insert_heading("ACRONYMS AND ABBREVIATIONS")
    for abbr, full_form in acronyms_data.items():
        p = abstract_para.insert_paragraph_before()
        # manual alignment hack
        p.add_run(f"{abbr}").bold = True
        p.add_run(f"\t- {full_form}")
        # Add a tab stop to align definitions
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls
        tab_stops = parse_xml(r'<w:tabs %s><w:tab w:val="left" w:pos="1440"/></w:tabs>' % nsdecls('w'))
        p.paragraph_format._element.get_or_add_pPr().append(tab_stops)

    insert_pb()

    doc.save(doc_path)
    print("Successfully built new Frontmatter: TOC, TOF, TOT, and Acronyms!")

if __name__ == '__main__':
    rebuild_frontmatter()