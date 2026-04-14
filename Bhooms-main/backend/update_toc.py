from docx import Document

def update_toc():
    doc_path = r'c:\Users\Harsha\Documents\muse-ai\muse-ai\Muse_AI_College_Project_Report_Final_Structured.docx'
    doc = Document(doc_path)
    
    # Update Table of Contents
    toc_start = -1
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().upper() == 'TABLE OF CONTENTS':
            toc_start = i
            break
            
    if toc_start != -1:
        # Find where to insert in TOC. 
        for i in range(toc_start + 1, min(toc_start + 50, len(doc.paragraphs))):
            text_lower = doc.paragraphs[i].text.lower()
            if 'system architecture overview' in text_lower or 'system design' in text_lower:
                # Add sub-item for the new feature inserted below this chapter in TOC
                new_p = doc.paragraphs[i+1].insert_paragraph_before("    • Application Flow and Custom Hero & Body API Generation")
                break
                
    # Update List of Figures
    lof_start = -1
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().upper() == 'LIST OF FIGURES':
            lof_start = i
            break
            
    if lof_start != -1:
        for i in range(lof_start + 1, min(lof_start + 30, len(doc.paragraphs))):
            text_lower = doc.paragraphs[i].text.lower()
            if 'architecture' in text_lower or 'design' in text_lower:
                doc.paragraphs[i+1].insert_paragraph_before("    • Figure: Muse AI Core Application and Custom Layout Pipeline")
                break
                
    doc.save(doc_path)
    print("Table of Contents and List of Figures updated successfully!")

if __name__ == '__main__':
    update_toc()